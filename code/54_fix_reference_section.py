"""
Reference-section APA-7 cleanup pass on Perron-SSWR-AI-2010_2015.docx.

Changes:
  1. Merge each split-across-two-paragraphs reference into a single paragraph
  2. Apply italics correctly to journal titles, volume numbers, working-paper
     and preprint titles
  3. Fill missing structural elements: editors, status markers, DOI/URL,
     version parentheticals, site names
  4. Restructure Desklib as APA-7 software citation
  5. Insert Grammarly (n.d.) and Daghestani (2018) placeholder entries in
     alphabetical order; both flagged with [verify ...] markers since the
     specific URL/source were not provided

Operates in place on Perron-SSWR-AI-2010_2015.docx.
"""

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from pathlib import Path

DOC = Path("/Users/beperron/Documents/GitHub/sswr-ai-authorship/Perron-SSWR-AI-2010_2015.docx")
d = Document(DOC)


def rewrite_in_place(first_para, second_para, segments):
    """Replace first_para's runs with the new segments. Delete second_para
    if not None. Each segment is (text, italic_bool)."""
    for r in list(first_para.runs):
        r._element.getparent().remove(r._element)
    for text, italic in segments:
        run = first_para.add_run(text)
        run.italic = italic
    if second_para is not None:
        second_para._element.getparent().remove(second_para._element)


def insert_after(prev_para, segments):
    new_p_el = OxmlElement("w:p")
    prev_para._p.addnext(new_p_el)
    new_p = Paragraph(new_p_el, prev_para._parent)
    new_p.style = prev_para.style
    for text, italic in segments:
        run = new_p.add_run(text)
        run.italic = italic
    return new_p


def find_para(starts_with):
    for p in d.paragraphs:
        if p.text.startswith(starts_with):
            return p
    return None


# ---------- 1. Capture (first, second) paragraph pairs for each ref ----------
# Indices were verified before edits; we use them to grab paragraph objects.
paras = list(d.paragraphs)

# Identify references by the prefix of the first paragraph
ref_pairs = {
    "Brynjolfsson":      ("Brynjolfsson, E.", True),
    "Dell'Acqua":        ("Dell'Acqua",        True),
    "Desklib":           ("Desklib. (2025)",   True),
    "Dugan":             ("Dugan, L.,",        True),
    "Gartenberg":        ("Gartenberg, C.,",   True),
    "Horta":             ("Horta, H., & Jung", True),
    "Hyland":            ("Hyland, K. (2005)", False),  # single-paragraph already
    "Liang":             ("Liang, W.,",        True),
    "Meyerson":          ("Meyerson, L. A.",   True),
    "Morley":            ("Morley, C. P.",     True),
    "Pangram":           ("Pangram Labs.",     True),
    "Perron_tenure":     ("Perron, B. E. (2026)", True),
    "Perron_et_al":      ("Perron, B. E., Victor", True),
    "Thai":              ("Thai, K.,",         True),
}


def get_ref_pair(prefix, has_second):
    """Find the first paragraph starting with `prefix` and (if has_second) the
    next non-empty paragraph after it that is part of the same reference."""
    parent = None
    first = None
    for i, p in enumerate(d.paragraphs):
        if p.text.startswith(prefix):
            first = p
            second = None
            if has_second and i + 1 < len(d.paragraphs):
                second = d.paragraphs[i + 1]
            return first, second
    return None, None


# ---------- 2. Rewrite each existing reference ----------
# Each ref is rewritten as a single paragraph with explicit italic runs.

# Brynjolfsson, Li, & Raymond (2023) — Working paper, italic title
first, second = get_ref_pair(*ref_pairs["Brynjolfsson"])
rewrite_in_place(first, second, [
    ("Brynjolfsson, E., Li, D., & Raymond, L. R. (2023). ", False),
    ("Generative AI at work", True),
    (" (Working Paper No. 31161). National Bureau of Economic Research. "
     "https://doi.org/10.3386/w31161", False),
])
print("[Brynjolfsson] merged + italicized title + DOI added")

# Dell'Acqua et al. (2023) — Working paper, italic title, add DOI
first, second = get_ref_pair(*ref_pairs["Dell'Acqua"])
rewrite_in_place(first, second, [
    ("Dell'Acqua, F., McFowland, E., III, Mollick, E. R., Lifshitz-Assaf, H., "
     "Kellogg, K., Rajendran, S., Krayer, L., Candelon, F., & Lakhani, K. R. (2023). ", False),
    ("Navigating the jagged technological frontier: Field experimental evidence "
     "of the effects of AI on knowledge worker productivity and quality", True),
    (" (Harvard Business School Working Paper No. 24-013). Harvard Business "
     "School. https://doi.org/10.2139/ssrn.4573321", False),
])
print("[Dell'Acqua] merged + italicized title + DOI added")

# Desklib (2025) — Software citation, restructure to APA-7 format
first, second = get_ref_pair(*ref_pairs["Desklib"])
rewrite_in_place(first, second, [
    ("Desklib. (2025). ", False),
    ("AI-Text-Detector-Academic", True),
    (" (Version 1.01) [Computer software]. Hugging Face. "
     "https://huggingface.co/desklib/ai-text-detector-academic-v1.01", False),
])
print("[Desklib] restructured as APA-7 software citation")

# Dugan et al. (2024) — Conference proceedings; add editors
first, second = get_ref_pair(*ref_pairs["Dugan"])
rewrite_in_place(first, second, [
    ("Dugan, L., Hwang, A., Trhlík, F., Ludan, J. M., Zhu, A., Xu, H., "
     "Ippolito, D., & Callison-Burch, C. (2024). RAID: A shared benchmark "
     "for robust evaluation of machine-generated text detectors. In L.-W. Ku, "
     "A. Martins, & V. Srikumar (Eds.), ", False),
    ("Proceedings of the 62nd Annual Meeting of the Association for "
     "Computational Linguistics (Volume 1: Long Papers)", True),
    (" (pp. 12463–12492). Association for Computational Linguistics. "
     "https://aclanthology.org/2024.acl-long.674", False),
])
print("[Dugan] merged + added editors")

# Gartenberg et al. (2026) — Italic journal name; preserve existing DOI
first, second = get_ref_pair(*ref_pairs["Gartenberg"])
rewrite_in_place(first, second, [
    ("Gartenberg, C., Hasan, S., Murray, A., & Pierce, L. (2026). More versus "
     "better: Artificial intelligence, incentives, and the emerging crisis "
     "in peer review. ", False),
    ("Organization Science", True),
    (". Advance online publication. https://doi.org/10.1287/orsc.2026.ed.v37.n3", False),
])
print("[Gartenberg] merged + journal italicized")

# Horta & Jung (2024) — Italicize journal + volume
first, second = get_ref_pair(*ref_pairs["Horta"])
rewrite_in_place(first, second, [
    ("Horta, H., & Jung, J. (2024). The crisis of peer review: Part of the "
     "evolution of science. ", False),
    ("Higher Education Quarterly", True),
    (", ", False),
    ("78", True),
    ("(4), Article e12511. https://doi.org/10.1111/hequ.12511", False),
])
print("[Horta & Jung] merged + journal/volume italicized")

# Hyland (2005) — already a single paragraph; rewrite for run consistency
first, _ = get_ref_pair(*ref_pairs["Hyland"])
rewrite_in_place(first, None, [
    ("Hyland, K. (2005). ", False),
    ("Metadiscourse: Exploring interaction in writing", True),
    (". Continuum.", False),
])
print("[Hyland] preserved (period moved outside italics)")

# Liang et al. (2023) — already mostly correct; rewrite for run consistency + ensure full italics
first, second = get_ref_pair(*ref_pairs["Liang"])
rewrite_in_place(first, second, [
    ("Liang, W., Yuksekgonul, M., Mao, Y., Wu, E., & Zou, J. (2023). GPT "
     "detectors are biased against non-native English writers. ", False),
    ("Patterns", True),
    (", ", False),
    ("4", True),
    ("(7), Article 100779. https://doi.org/10.1016/j.patter.2023.100779", False),
])
print("[Liang] merged + verified italics")

# Meyerson et al. (2025) — Italicize journal + volume
first, second = get_ref_pair(*ref_pairs["Meyerson"])
rewrite_in_place(first, second, [
    ("Meyerson, L. A., Suzzi-Simmons, A., & Simberloff, D. (2025). Quantifying "
     "reviewer declines in scientific publishing: Twenty-one years of data "
     "from Biological Invasions 2002–2024. ", False),
    ("Biological Invasions", True),
    (", ", False),
    ("27", True),
    ("(10), Article 223. https://doi.org/10.1007/s10530-025-03679-1", False),
])
print("[Meyerson] merged + journal/volume italicized")

# Morley et al. (2025) — Italicize journal + volume
first, second = get_ref_pair(*ref_pairs["Morley"])
rewrite_in_place(first, second, [
    ("Morley, C. P., Prunuske, J., Phillips, J. P., Wendling, A. L., "
     "Heidelbaugh, J. J., & Grammer, S. (2025). Reviewer engagement trends "
     "at a journal: Cause for concern. ", False),
    ("PRiMER", True),
    (", ", False),
    ("9", True),
    (", Article 59. https://doi.org/10.22454/PRiMER.2025.831615", False),
])
print("[Morley] merged + journal/volume italicized")

# Pangram Labs (2025) — Add site name
first, second = get_ref_pair(*ref_pairs["Pangram"])
rewrite_in_place(first, second, [
    ("Pangram Labs. (2025, November). ", False),
    ("Pangram predicts 21% of ICLR reviews are AI-generated", True),
    (". Pangram Blog. "
     "https://www.pangram.com/blog/pangram-predicts-21-of-iclr-reviews-are-ai-generated", False),
])
print("[Pangram Labs] merged + site name added")

# Perron (2026) [tenure paper] — Italicize title
first, second = get_ref_pair(*ref_pairs["Perron_tenure"])
rewrite_in_place(first, second, [
    ("Perron, B. E. (2026). ", False),
    ("Tenure and promotion in the age of AI: Why academia needs to start "
     "talking now", True),
    (" [Working paper]. SSRN. https://ssrn.com/abstract=6640399", False),
])
print("[Perron tenure] merged + title italicized")

# Perron, Victor, & Qi (2026) — Italicize journal, add 'Advance online publication'
first, second = get_ref_pair(*ref_pairs["Perron_et_al"])
rewrite_in_place(first, second, [
    ("Perron, B. E., Victor, B. G., & Qi, Z. (2026). AI-assisted curation "
     "of conference scholarship: Compiling, structuring, and analyzing two "
     "decades of presentations at the Society for Social Work and Research "
     "(2005–2026). ", False),
    ("Journal of the Society for Social Work and Research", True),
    (". Advance online publication.", False),
])
print("[Perron, Victor, & Qi] merged + journal italicized + status marker added")

# Thai et al. (2025) — Italicize preprint title; switch to canonical DOI form
first, second = get_ref_pair(*ref_pairs["Thai"])
rewrite_in_place(first, second, [
    ("Thai, K., Emi, B., Masrour, E., & Iyyer, M. (2025). ", False),
    ("EditLens: Quantifying the extent of AI editing in text", True),
    (" (arXiv:2510.03154) [Preprint]. arXiv. "
     "https://doi.org/10.48550/arXiv.2510.03154", False),
])
print("[Thai] merged + title italicized + DOI form")


# ---------- 3. Insert Grammarly (n.d.) and Daghestani (2018) placeholders ----------
# Daghestani goes between Brynjolfsson and Dell'Acqua
brynj = find_para("Brynjolfsson, E.")
if brynj is not None:
    insert_after(brynj, [
        ("Daghestani, A. (2018). Grammarly Series A funding and active-user "
         "growth [verify primary source]. ", False),
        ("[Source to verify — likely TechCrunch / Forbes / Crunchbase coverage "
         "of the December 2017 General Catalyst Series A round].", True),
    ])
    print("[Daghestani 2018] placeholder inserted (verify source)")

# Grammarly goes between Gartenberg and Horta
gart = find_para("Gartenberg, C.,")
if gart is not None:
    insert_after(gart, [
        ("Grammarly. (n.d.). ", False),
        ("About Grammarly", True),
        (". Retrieved May 5, 2026, from https://www.grammarly.com/about "
         "[verify final URL: corporate ", False),
        ("about", True),
        (" page or ", False),
        ("press kit", True),
        ("]", False),
    ])
    print("[Grammarly n.d.] placeholder inserted (verify URL)")


# ---------- save ----------
d.save(DOC)
print(f"\nSaved -> {DOC}")
