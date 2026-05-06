"""
Apply Hassan/Gartenberg-code primary writing-quality results to
reports/FinalReport_edited.docx, updating:

  1. Writing-Quality Measures paragraph (Hassan/Gartenberg primary, textstat secondary)
  2. Table 5 (writing-quality ITS step coefficients)
  3. Writing-Quality Trajectory body paragraph (Pearson r values)
  4. Embedded image bytes for Figures 1, 2, 3

Reads:
  results/writing_quality_its_hassan.json
  results/figures/yearly_three_detector_lines.png
  results/figures/fig_writing_quality_trend.png
  results/figures/fig_writing_quality_forest.png

Writes (in place):
  reports/FinalReport_edited.docx
"""

import json, shutil, zipfile, io
from pathlib import Path
from docx import Document
from PIL import Image

ROOT = Path("/Users/beperron/Desktop/AI-SSWR")
DOCX = ROOT / "reports" / "FinalReport_edited.docx"
JSON = ROOT / "results" / "writing_quality_its_hassan.json"
FIGS = ROOT / "results" / "figures"

R = json.loads(JSON.read_text())
its = R["its"]

# ---- 1+3: text edits ---------------------------------------------------
doc = Document(str(DOCX))

# --- paragraph 43: Writing-Quality Measures ---
new_p43 = (
    "I selected seven of nine writing-quality metrics from Gartenberg et al. "
    "(2026). The seven retained metrics are Flesch Reading Ease, Flesch-Kincaid "
    "Grade Level, Gunning FOG Index, SMOG Index, nominalization rate, "
    "passive-voice rate, and hedging rate (Hyland, 2005). The primary "
    "computation used the reference Python code from the Gartenberg study "
    "(included in the project repository as code/44_hassan_analyzer.py); "
    "running Gartenberg et al.'s own metric implementation on the SSWR "
    "corpus yields a stricter form of conceptual replication than a "
    "from-scratch reimplementation. As an independent robustness check, all "
    "seven metrics were re-computed in a self-contained reimplementation "
    "using the textstat library for readability and spaCy part-of-speech "
    "and dependency parsing for style metrics; the two pipelines produce "
    "identically signed steps at the partial → full exposure boundary and "
    "correlate Pearson r = 0.74–0.89 on readability. Jargon and specificity "
    "metrics were omitted because their operationalization relies on "
    "field-specific reference vocabularies (management/organization-science "
    "terminology and economics theoretical markers) that do not transfer to "
    "a social-work corpus."
)
p43 = doc.paragraphs[43]
# Replace by clearing all runs and writing one run; keep paragraph style
# (preserves the 0.5" indent etc. assigned at the paragraph level).
for run in p43.runs:
    run.text = ""
p43.runs[0].text = new_p43
print("[OK] Updated paragraph 43 (Writing-Quality Measures)")

# --- paragraph 64: Writing-Quality Trajectory body ---
fre_r  = R["pearson_r_vs_editlens"]["Flesch Reading Ease"]["r"]
pas_r  = R["pearson_r_vs_editlens"]["Passive-voice rate"]["r"]
def fmtr(v):  # use Unicode minus
    return f"{v:.2f}".replace("-", "−")

new_p64 = (
    f"Table 5 reports the interrupted time-series step at the partial → full "
    f"exposure boundary for each of the seven Gartenberg-aligned "
    f"writing-quality metrics, computed from the Gartenberg-study reference "
    f"code applied to the SSWR corpus. Six metrics produced a step at "
    f"p < .0001 and the seventh (hedging rate) at p = .016. The Pearson "
    f"correlation between Flesch Reading Ease and EditLens RoBERTa-large "
    f"score on the full corpus is r = {fmtr(fre_r)} (p < .0001), in the same "
    f"direction as Gartenberg et al. (2026) reported (r = −0.4 with Pangram). "
    f"Passive voice shows the strongest negative correlation among style "
    f"metrics (r = {fmtr(pas_r)})."
)
p64 = doc.paragraphs[64]
for run in p64.runs:
    run.text = ""
p64.runs[0].text = new_p64
print("[OK] Updated paragraph 64 (Writing-Quality Trajectory)")

# --- paragraph 73: Replication of Gartenberg (Pearson r = -0.31) ---
p73 = doc.paragraphs[73]
new_p73 = p73.text.replace("r = −0.31", f"r = {fre_r:.2f}".replace("-", "−"))
if new_p73 != p73.text:
    for run in p73.runs:
        run.text = ""
    p73.runs[0].text = new_p73
    print("[OK] Updated paragraph 73 Pearson r")
else:
    print("[OK] Paragraph 73 r value already matches")

# ---- 2: Table 5 (= python-docx Tables[4]) -----------------------------
def fmt(label):
    r = its[label]
    s, lo, hi = r["step"], r["ci_lo"], r["ci_hi"]
    # 2 decimals for |s| >= 0.1, otherwise 4 decimals to preserve small-rate precision
    if abs(s) >= 0.1:
        nd = 2
    else:
        nd = 4
    fmtspec = f"+.{nd}f"
    s_str  = format(s,  fmtspec).replace("-", "−")
    lo_str = format(lo, fmtspec).replace("-", "−")
    hi_str = format(hi, fmtspec).replace("-", "−")
    p = r["p"]
    p_str = "< .0001" if p < .0001 else f"= {p:.3f}".replace("0.", ".")
    return s_str, f"[{lo_str}, {hi_str}]", p_str

# Table 5 = Tables[4]
t5 = doc.tables[4]
labels = [
    "Flesch Reading Ease",
    "Flesch-Kincaid Grade Level",
    "Gunning FOG Index",
    "SMOG Index",
    "Nominalization rate",
    "Passive-voice rate",
    "Hedging rate (Hyland)",
]
for i, label in enumerate(labels, start=1):
    row = t5.rows[i]
    s, ci, p = fmt(label)
    # cell 0 already labeled; update cells 1, 2, 3
    def set_cell(cell, txt):
        # clear existing then write, preserving paragraph alignment
        for para in cell.paragraphs:
            for run in para.runs:
                run.text = ""
            if para.runs:
                para.runs[0].text = txt
            else:
                para.add_run(txt)
            break
    set_cell(row.cells[1], s)
    set_cell(row.cells[2], ci)
    set_cell(row.cells[3], p)
    print(f"[OK] Table 5 row '{label}': {s}  {ci}  {p}")

# ---- Save text edits to a temp file -----------------------------------
TMP = ROOT / "reports" / "_FinalReport_tmp.docx"
doc.save(str(TMP))
print(f"\nSaved text-edited intermediate -> {TMP}")

# ---- 4: Replace embedded images ---------------------------------------
# Convert each new PNG to high-quality JPG (white BG) for in-docx embedding;
# the standalone PNG/SVG copies in results/figures/ are kept untouched.
fig_map = {
    "word/media/image1.jpg": FIGS / "yearly_three_detector_lines.png",
    "word/media/image2.jpg": FIGS / "fig_writing_quality_trend.png",
    "word/media/image3.jpg": FIGS / "fig_writing_quality_forest.png",
}

new_jpg_bytes = {}
for media_path, src_png in fig_map.items():
    img = Image.open(src_png)
    if img.mode in ("RGBA","LA"):
        bg = Image.new("RGB", img.size, (255,255,255))
        bg.paste(img, mask=img.split()[-1])
        img = bg
    else:
        img = img.convert("RGB")
    buf = io.BytesIO()
    img.save(buf, format="JPEG", quality=95, optimize=True, subsampling=0)
    new_jpg_bytes[media_path] = buf.getvalue()
    print(f"[OK] Converted {src_png.name} -> {len(buf.getvalue()):,} bytes JPG ({media_path})")

# Rebuild the docx zip with replaced media
OUT_TMP = ROOT / "reports" / "_FinalReport_imgs.docx"
with zipfile.ZipFile(TMP, "r") as zin, zipfile.ZipFile(OUT_TMP, "w", zipfile.ZIP_DEFLATED) as zout:
    for item in zin.infolist():
        data = zin.read(item.filename)
        if item.filename in new_jpg_bytes:
            data = new_jpg_bytes[item.filename]
            print(f"[OK] Replaced {item.filename}")
        zout.writestr(item, data)

# Move into place
shutil.move(str(OUT_TMP), str(DOCX))
TMP.unlink()
print(f"\nFinal -> {DOCX}")
