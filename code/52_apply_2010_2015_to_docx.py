"""
Apply all 2010-2015 calibration updates to Perron-SSWR-AI-FINAL.docx.
Saves to a NEW file (Perron-SSWR-AI-2010_2015.docx) so the original is preserved.

Changes made (in order):
  1.  Drop methodology sentence in Limitations (Para [0115])
  2.  Calibration string replacements throughout
  3.  P95 thresholds in Methods (Para [0074])
  4.  Abstract H1 numbers (Para [0039])
  5.  Body H1 range sentence (Para [0086]) — fixes lowest-detector misstatement
  6.  Table 1 (yearly) cells
  7.  H3 country section (Para [0094]) — numbers + "H4 is supported" -> "H3 is supported"
  8.  H4 paragraph rewrite (Para [0096])
  9.  Cohen's kappa (Para [0101])
  10. Figure 2 caption rewrite (Para [0181])
  11. Table 0 row 3 (8 years -> 6 years)
  12. Prior-subs section (Paras [0090], [0091], [0165]) and Table 2 cells
  13. Pre-LLM drift mention in Discussion (Para [0109]) cleanup if needed
"""

from docx import Document
from pathlib import Path
import shutil

REPO = Path("/Users/beperron/Documents/GitHub/sswr-ai-authorship")
SRC  = REPO / "Perron-SSWR-AI-FINAL.docx"
DST  = REPO / "Perron-SSWR-AI-2010_2015.docx"

shutil.copy(SRC, DST)
d = Document(DST)


# -------- helpers ---------------------------------------------------------
def replace_in_paragraph(p, old, new):
    """Replace `old` substring with `new` inside a paragraph.
    Tries to preserve run formatting if the substring lives in a single run.
    Falls back to flattening to a single run if it spans runs.
    Returns True if a replacement was made."""
    full_text = "".join(r.text for r in p.runs)
    if old not in full_text:
        return False
    # Single-run case first
    for run in p.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            return True
    # Spans runs — flatten
    runs = list(p.runs)
    if not runs:
        p.add_run(full_text.replace(old, new))
        return True
    # Keep first run's formatting; put the entire new text in it; clear the rest
    runs[0].text = full_text.replace(old, new)
    for r in runs[1:]:
        r._element.getparent().remove(r._element)
    return True


def set_paragraph_text(p, new_text):
    """Wholesale replace paragraph text. First-run formatting is kept; sub-run
    formatting (e.g., italic on a span) is lost."""
    runs = list(p.runs)
    if not runs:
        p.add_run(new_text)
        return
    runs[0].text = new_text
    for r in runs[1:]:
        r._element.getparent().remove(r._element)


def find_paragraph_index_containing(text_fragment, start_at=0):
    """Return paragraph index whose text contains `text_fragment`, or -1."""
    for i, p in enumerate(d.paragraphs):
        if i < start_at:
            continue
        if text_fragment in p.text:
            return i
    return -1


def cell_set(table_idx, row_idx, col_idx, new_text):
    cell = d.tables[table_idx].rows[row_idx].cells[col_idx]
    if cell.paragraphs:
        set_paragraph_text(cell.paragraphs[0], new_text)
        for extra in cell.paragraphs[1:]:
            extra._element.getparent().remove(extra._element)
    else:
        cell.text = new_text


# -------- step 1: drop the methodology sentence in Limitations -----------
i = find_paragraph_index_containing("H3 yielded a result opposite to the predicted direction")
if i >= 0:
    new_limit = ("The full-exposure period contains only two conference cycles, so "
                 "level-change estimates are stable but post-2026 trajectory cannot be "
                 "projected. Two Gartenberg metrics (Jargon and Specificity) were "
                 "omitted because their operationalization relies on field-specific "
                 "reference vocabularies (see Methods §Writing-Quality Measures).")
    set_paragraph_text(d.paragraphs[i], new_limit)
    print(f"[1] Dropped methodology sentence (Para [{i:04d}])")


# -------- step 2: calibration string replacements -----------------------
# Apply across all paragraphs and table cells
# Note: only replace where the surrounding context confirms it's about
# calibration, not about the analytic/exposure window. The phrase "2010–2017"
# appears only in calibration contexts in this document — verified.
calibration_replacements = [
    ("2010–2017", "2010–2015"),
    ("2010-2017", "2010-2015"),
    ("post-2017 abstract", "post-2015 abstract"),
    ("post-2017 score distribution", "post-2015 score distribution"),
]
n_para = 0
for p in d.paragraphs:
    for old, new in calibration_replacements:
        if replace_in_paragraph(p, old, new):
            n_para += 1
n_cell = 0
for table in d.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for old, new in calibration_replacements:
                    if replace_in_paragraph(p, old, new):
                        n_cell += 1
print(f"[2] Calibration string replacements: {n_para} in body, {n_cell} in tables")


# -------- step 3: P95 thresholds in Methods (Para [0074]) ---------------
# OLD: "EditLens RoBERTa-large = 0.109 (single-window) and 0.207 (multi-window aggregate);
#       EditLens Llama-3.2-3B = 0.062; and Desklib academic = 0.999"
i = find_paragraph_index_containing("EditLens RoBERTa-large = 0.109")
if i >= 0:
    p = d.paragraphs[i]
    # P95 single
    replace_in_paragraph(p, "EditLens RoBERTa-large = 0.109 (single-window) and 0.207 (multi-window aggregate)",
                         "EditLens RoBERTa-large = 0.105 (single-window) and 0.201 (multi-window aggregate)")
    replace_in_paragraph(p, "EditLens Llama-3.2-3B = 0.062",
                         "EditLens Llama-3.2-3B = 0.056")
    # desklib stays at 0.999, no change needed
    print(f"[3] P95 thresholds updated (Para [{i:04d}])")


# -------- step 4: Abstract H1 numbers (Para [0039]) ---------------------
i = find_paragraph_index_containing("rose from a stable 5.0% across 2010")
if i >= 0:
    p = d.paragraphs[i]
    replace_in_paragraph(p, "to 18.7% in 2024, 32.2% in 2025, and 56.1% in 2026",
                         "to 19.6% in 2024, 33.3% in 2025, and 57.0% in 2026")
    print(f"[4] Abstract H1 numbers updated (Para [{i:04d}])")


# -------- step 5: Body H1 range (Para [0086]) ---------------------------
i = find_paragraph_index_containing("By conference 2026 the percentage of abstracts above threshold ranges from")
if i >= 0:
    p = d.paragraphs[i]
    replace_in_paragraph(
        p,
        "ranges from 59.3% (Desklib academic) to 82.9% (EditLens Llama-3.2-3B)",
        "ranges from 57.0% (EditLens RoBERTa-large) to 84.3% (EditLens Llama-3.2-3B)")
    print(f"[5] Body H1 range updated (Para [{i:04d}])")


# -------- step 6: Table 1 (yearly) cells --------------------------------
# Table 1 has rows R1..R17 = years 2010..2026, cols 2/3/4 = R/Llama/Desklib
yearly_new = [
    ("2010",   "3.3",  "2.5",  "5.1"),
    ("2011",   "3.8",  "3.2",  "4.6"),
    ("2012",   "4.6",  "3.0",  "5.1"),
    ("2013",   "6.4",  "6.9",  "5.0"),
    ("2014",   "4.2",  "6.0",  "4.1"),
    ("2015",   "6.4",  "6.5",  "5.7"),
    ("2016",   "5.9",  "7.6",  "4.8"),
    ("2017",   "6.4",  "9.8",  "4.4"),
    ("2018",   "8.8",  "10.9", "5.6"),
    ("2019",   "7.9",  "14.0", "6.3"),
    ("2020",   "8.5",  "18.8", "6.2"),
    ("2021",   "9.8",  "24.9", "7.9"),
    ("2022",  "12.9",  "31.9", "8.2"),
    ("2023",  "14.9",  "36.7", "9.7"),
    ("2024",  "19.6",  "50.8","14.9"),
    ("2025",  "33.3",  "66.4","31.3"),
    ("2026",  "57.0",  "84.3","59.1"),
]
for ri, (yr, e, l, k) in enumerate(yearly_new, start=1):
    cell_set(1, ri, 2, e)   # EditLens RoBERTa-large
    cell_set(1, ri, 3, l)   # EditLens Llama
    cell_set(1, ri, 4, k)   # Desklib
print(f"[6] Table 1 yearly cells updated (51 cells)")


# -------- step 7: H3 country (Para [0094]) ------------------------------
i = find_paragraph_index_containing("first-author country grouped as US")
if i >= 0:
    p = d.paragraphs[i]
    repls = [
        ("no-exposure rate 7.8%, n = 14,916; full-exposure rate 43.9%, n = 3,180",
         "no-exposure rate 8.4%, n = 14,916; full-exposure rate 44.8%, n = 3,180"),
        ("8.5%, n = 612 → 52.9%, n = 121",
         "9.5%, n = 612 → 53.7%, n = 121"),
        ("3.2%, n = 818 → 50.0%, n = 362",
         "3.3%, n = 818 → 51.7%, n = 362"),
        ("+36.1 percentage points (US), +44.4 (Other Anglophone), and +46.8 (Non-Anglophone)",
         "+36.4 percentage points (US), +44.2 (Other Anglophone), and +48.4 (Non-Anglophone)"),
        ("Other-Anglophone-by-full coefficient of +8.3 percentage points (p = .005, 95% CI [+2.5, +14.2])",
         "Other-Anglophone-by-full coefficient of +7.8 percentage points (p = .002, 95% CI [+2.8, +12.8])"),
        ("Non-Anglophone-by-full coefficient of +10.8 percentage points (p < .0001, 95% CI [+9.0, +12.6])",
         "Non-Anglophone-by-full coefficient of +11.9 percentage points (p < .0001, 95% CI [+9.4, +14.5])"),
        ("H4 is supported", "H3 is supported"),
    ]
    for old, new in repls:
        replace_in_paragraph(p, old, new)
    print(f"[7] H3 country section updated (Para [{i:04d}])")


# -------- step 8: H4 paragraph rewrite (Para [0096]) --------------------
i = find_paragraph_index_containing("Within the 2010–2015 calibration baseline, mean EditLens RoBERTa-large")
if i < 0:
    # in case step 2 replaced 2010-2017 to 2010-2015 inside this paragraph
    i = find_paragraph_index_containing("calibration baseline, mean EditLens RoBERTa-large")
if i >= 0:
    new_h4 = (
        "Within the 2010–2015 calibration baseline, mean EditLens RoBERTa-large "
        "scores moved from 0.0354 in 2010 to 0.0381 in 2015, an accumulated change "
        "of approximately 0.3 percentage points across the six-year window. The "
        "academic-tuned Desklib detector moved from 0.624 in 2010 to 0.662 in 2015, "
        "an accumulated change of approximately 3.8 percentage points. Year-clustered "
        "abstract-level regressions of detector score on year and log word count "
        "tested whether this within-baseline drift was statistically distinguishable "
        "from zero. On EditLens RoBERTa-large the year coefficient was +0.0004 "
        "(p = .18), supporting H4 on the primary detector. On academic-tuned Desklib "
        "the year coefficient was +0.0103 (p < .0001), rejecting H4 on this secondary "
        "detector. The Desklib drift is consistent with the variant's greater "
        "sensitivity to subtle stylistic shifts, including the Grammarly-era editing "
        "tools described in the discussion. The practical magnitude of the Desklib "
        "drift, however, is small relative to the post-LLM signal: the full-exposure "
        "level change at the 2025 conference cycle is +20.9 percentage points on "
        "Desklib, more than five times larger than the accumulated within-baseline "
        "drift. On the primary detector the full-exposure step (+17.8 percentage "
        "points) sits against a baseline that is statistically flat. The post-LLM "
        "step is therefore not plausibly attributable to continuation of pre-LLM "
        "stylistic drift on either detector."
    )
    set_paragraph_text(d.paragraphs[i], new_h4)
    print(f"[8] H4 paragraph rewritten (Para [{i:04d}])")


# -------- step 9: Cohen's kappa (Para [0101]) ---------------------------
i = find_paragraph_index_containing("Cohen’s κ on binary classifications is 0.41 (Desklib academic vs EditLens RoBERTa-large)")
if i < 0:
    i = find_paragraph_index_containing("Cohen's κ on binary classifications is 0.41 (Desklib academic vs EditLens RoBERTa-large)")
if i >= 0:
    p = d.paragraphs[i]
    # Old: "0.41 (Desklib academic vs EditLens RoBERTa-large), 0.41 (Desklib ... Llama), and 0.41 (RoBERTa ... Llama)"
    # New: "0.41 (Desklib academic vs EditLens RoBERTa-large), 0.38 (Desklib ... Llama), and 0.39 (RoBERTa ... Llama)"
    replace_in_paragraph(p,
        "0.41 (Desklib academic vs EditLens Llama-3.2-3B), and 0.41 (EditLens RoBERTa-large vs Llama-3.2-3B)",
        "0.38 (Desklib academic vs EditLens Llama-3.2-3B), and 0.39 (EditLens RoBERTa-large vs Llama-3.2-3B)")
    print(f"[9] Cohen's kappa updated (Para [{i:04d}])")


# -------- step 10: Figure 2 caption (Para [0181]) -----------------------
i = find_paragraph_index_containing("Mean Flesch Reading Ease per SSWR conference cycle, standardized to the")
if i >= 0:
    new_fig2 = (
        "Note. Mean Flesch Reading Ease per SSWR conference cycle, standardized "
        "to the 2010–2015 baseline (mean = 0, SD = 1). Shaded band shows the 95% "
        "confidence interval of the cycle mean. The 2010–2016 cycles are omitted "
        "from the plot for clarity (each within −0.07 to +0.05 SD of baseline) and "
        "indicated by the “//” axis break. Era shading indicates the no-exposure "
        "(gray), partial-exposure (light peach, conf 2024), and full-exposure "
        "(peach, conf 2025–2026) periods."
    )
    set_paragraph_text(d.paragraphs[i], new_fig2)
    print(f"[10] Figure 2 caption rewritten (Para [{i:04d}])")


# -------- step 11: Table 0 R3 (8 years -> 6 years) ----------------------
# After step 2 the cell already says "Conferences 2010–2015 (8 years)" — needs (6 years).
cell = d.tables[0].rows[3].cells[2]
for p in cell.paragraphs:
    replace_in_paragraph(p, "(8 years)", "(6 years)")
print("[11] Table 0 R3: '8 years' -> '6 years'")


# -------- step 12: Prior-subs (Para [0090], [0091], [0165] + Table 2) ---
# Para [0090]: rates and changes per bucket
i = find_paragraph_index_containing("New first authors rose 40.2 percentage points")
if i >= 0:
    p = d.paragraphs[i]
    # NEW (Method G):
    #   New (5,381):         no-exp 7.76%, full 50.12%, change +42.4 pp
    #   Early (4,854):       no-exp 7.73%, full 47.63%, change +39.9 pp
    #   Established (11,332): no-exp 8.60%, full 43.51%, change +34.9 pp
    new_text = (
        "Table 3 reports the rate of above-threshold flagging in each bucket and "
        "exposure period. All three buckets rose sharply from no exposure to full "
        "exposure, but the rise was steepest among authors with the fewest prior "
        "submissions. New first authors rose 42.4 percentage points (7.8% to 50.1%); "
        "Early-career first authors rose 39.9 percentage points (7.7% to 47.6%); and "
        "Established first authors rose 34.9 percentage points (8.6% to 43.5%)."
    )
    set_paragraph_text(p, new_text)
    print(f"[12a] Para [0090] prior-subs rates updated")

# Para [0091]: DiD coefficients
i = find_paragraph_index_containing("New first authors gained 6.12 percentage points more than Established")
if i >= 0:
    p = d.paragraphs[i]
    repls = [
        ("New first authors gained 6.12 percentage points more than Established first authors (95% CI [2.86 - 9.38], p = .0002)",
         "New first authors gained 7.45 percentage points more than Established first authors (95% CI [+6.47, +8.44], p < .0001)"),
        ("Early-career first authors gained 5.13 percentage points more (95% CI [3.49 - 6.77], p < .0001)",
         "Early-career first authors gained 4.99 percentage points more (95% CI [+3.03, +6.96], p < .0001)"),
        ("New and Early-career authors gained roughly 8 to 9 percentage points more than Established authors",
         "New and Early-career authors gained roughly 10 percentage points more than Established authors"),
        ("the corresponding gaps were 6 to 7 percentage points",
         "the corresponding gaps were 7.5 to 8.4 percentage points"),
        ("all four secondary-detector estimates significant at p < .0001",
         "all four secondary-detector estimates significant at p < .001"),
    ]
    for old, new in repls:
        replace_in_paragraph(p, old, new)
    print(f"[12b] Para [0091] prior-subs DiD updated")

# Table 2 (Prior-Submissions): R1=New, R2=Early, R3=Established
prior_subs_cells = [
    # row,  no_exp, partial, full,  change
    (1, "7.8",  "19.1", "50.1", "+42.4"),  # New
    (2, "7.7",  "27.2", "47.6", "+39.9"),  # Early-career
    (3, "8.6",  "17.3", "43.5", "+34.9"),  # Established
]
for ri, no_exp, part, full, chg in prior_subs_cells:
    cell_set(2, ri, 1, no_exp)
    cell_set(2, ri, 2, part)
    cell_set(2, ri, 3, full)
    cell_set(2, ri, 4, chg)
print("[12c] Table 2 (prior-subs) cells updated")

# Para [0165] Table 3 caption — bucket totals
i = find_paragraph_index_containing("New (0 prior submissions) n = 6,514")
if i >= 0:
    p = d.paragraphs[i]
    replace_in_paragraph(p,
        "New (0 prior submissions) n = 6,514; Early-career (1–2 prior submissions) n = 6,294; Established (3 or more prior submissions) n = 8,759",
        "New (0 prior submissions) n = 5,381; Early-career (1–2 prior submissions) n = 4,854; Established (3 or more prior submissions) n = 11,332")
    print(f"[12d] Table 3 caption bucket totals updated")


# -------- save ----------------------------------------------------------
d.save(DST)
print(f"\nSaved -> {DST}")
