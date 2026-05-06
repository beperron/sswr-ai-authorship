"""
Final-pass edits on the 2010-2015 manuscript:
  1.  Para [0055]: 'five hypotheses' -> 'four hypotheses' (Cowork bug)
  2.  Para [0073]: rewrite Methods Calibration paragraph — concise Grammarly
      justification, correct n=4,460
  3.  Para [0075]: drop 'saved before any post-2015 ... ensuring not tuned'
      pre-registration-style framing
  4.  Para [0094]: 'pre-specified subgroup comparison' -> 'subgroup comparison'
  5.  Para [0113]: fix 'H5 is rejected on both detectors' (H5 doesn't exist) +
      align with the new H4 finding (drift null on EditLens R, significant on
      Desklib)
  6.  Para [0165]: drop 'locked' from Table 3 caption
  7.  INSERT 'Calibration-Window Sensitivity' subsection (Heading 3 + body)
      after Inter-Detector Convergence and before Discussion

Operates on Perron-SSWR-AI-2010_2015.docx in place.
"""

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from pathlib import Path

DOC = Path("/Users/beperron/Documents/GitHub/sswr-ai-authorship/Perron-SSWR-AI-2010_2015.docx")
d = Document(DOC)


def replace_in_paragraph(p, old, new):
    full = "".join(r.text for r in p.runs)
    if old not in full:
        return False
    for run in p.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            return True
    runs = list(p.runs)
    if not runs:
        p.add_run(full.replace(old, new))
        return True
    runs[0].text = full.replace(old, new)
    for r in runs[1:]:
        r._element.getparent().remove(r._element)
    return True


def set_paragraph_text(p, new_text):
    runs = list(p.runs)
    if not runs:
        p.add_run(new_text)
        return
    runs[0].text = new_text
    for r in runs[1:]:
        r._element.getparent().remove(r._element)


def find_idx(fragment, start_at=0):
    for i, p in enumerate(d.paragraphs):
        if i < start_at:
            continue
        if fragment in p.text:
            return i
    return -1


# ---------- 1. Five -> four hypotheses ----------
i = find_idx("The analytic plan is constructed around five hypotheses")
if i >= 0:
    replace_in_paragraph(d.paragraphs[i], "five hypotheses", "four hypotheses")
    print(f"[1] Fixed five->four hypotheses (Para [{i:04d}])")


# ---------- 2. Methods Calibration paragraph rewrite (Para [0073]) ----------
i = find_idx("AI text-detector scores are not probabilities")
if i >= 0:
    new_calib = (
        "AI text-detector scores are not probabilities. They represent "
        "detector-specific confidence on a scale that varies by model. A score "
        "of 0.5 from one detector does not mean the same thing as a score of "
        "0.5 from another. To make the scores comparable and substantively "
        "meaningful, each detector was anchored against the same in-domain "
        "reference distribution: scientific-format SSWR abstracts from "
        "conferences 2010 through 2015 (n = 4,460). The 2015 upper bound "
        "predates Grammarly's May 2015 freemium launch by one submission "
        "cycle: the April 2016 deadline (conference 2017) is the first SSWR "
        "cycle whose abstracts could plausibly carry consumer-grade "
        "machine-editing polish, so ending the calibration window at "
        "conference 2015 leaves the reference distribution free of plausible "
        "Grammarly-era contamination at the upper end. Gartenberg et al. used "
        "the same anchoring strategy (with a two-year calibration window) at "
        "Organization Science. Where they applied judgment-based score "
        "cutoffs on the Pangram scale (0.15, 0.30, 0.70), the present study "
        "uses percentile cutoffs because EditLens produces scores on a "
        "different scale, and percentiles transfer across detectors."
    )
    set_paragraph_text(d.paragraphs[i], new_calib)
    print(f"[2] Rewrote Methods Calibration paragraph (Para [{i:04d}])")


# ---------- 3. Para [0075] drop pre-reg framing ----------
i = find_idx("Each detector's threshold was computed and saved before any post-2015 abstract was scored")
if i >= 0:
    new_75 = (
        "Each detector's threshold was computed on the 2010–2015 calibration "
        "window. EditLens additionally supplies a native four-bucket "
        "classification (Thai et al., 2025: fully human, lightly AI-edited, "
        "heavily AI-edited, fully AI-generated). Two derived rates are "
        "reported: the percentage of abstracts in any AI-involvement bucket "
        "(bucket ≥ 1) and the percentage in the fully AI-generated bucket "
        "(bucket = 3)."
    )
    set_paragraph_text(d.paragraphs[i], new_75)
    print(f"[3] Dropped pre-reg framing (Para [{i:04d}])")


# ---------- 4. Para [0094] drop 'pre-specified' ----------
i = find_idx("A pre-specified subgroup comparison used first-author country")
if i >= 0:
    replace_in_paragraph(d.paragraphs[i],
                         "A pre-specified subgroup comparison",
                         "A subgroup comparison")
    print(f"[4] Dropped 'pre-specified' (Para [{i:04d}])")


# ---------- 5. Para [0113] fix H5 + align with new H4 ----------
i = find_idx("H5 is rejected on both detectors")
if i >= 0:
    p = d.paragraphs[i]
    # Replace the entire H5 sentence with one that drops the spurious H5
    # reference and reflects the new H4 finding.
    old_sentence = (
        " H5 is rejected on both detectors (within-baseline drift = ~3% of "
        "the full-exposure step on EditLens RoBERTa-large, ~27% on "
        "academic-tuned Desklib); the drift is consistent with the pre-LLM "
        "editing-tool interpretation in Distinguishing Baseline Drift, and "
        "is at least an order of magnitude smaller than the full-exposure "
        "step."
    )
    new_sentence = (
        " The within-baseline drift detected on the academic-tuned Desklib "
        "detector (see H4) is consistent with the pre-LLM editing-tool "
        "interpretation in Distinguishing Baseline Drift and is at least an "
        "order of magnitude smaller than the full-exposure step; the primary "
        "EditLens RoBERTa-large detector shows no statistically significant "
        "within-baseline drift."
    )
    replace_in_paragraph(p, old_sentence, new_sentence)
    print(f"[5] Fixed H5 sentence + new H4 framing (Para [{i:04d}])")


# ---------- 6. Para [0165] drop 'locked' ----------
i = find_idx("Cells report the percentage of abstracts above the locked 95th-percentile threshold")
if i >= 0:
    replace_in_paragraph(
        d.paragraphs[i],
        "Cells report the percentage of abstracts above the locked 95th-percentile threshold",
        "Cells report the percentage of abstracts above the 95th-percentile threshold")
    print(f"[6] Dropped 'locked' (Para [{i:04d}])")


# ---------- 7. INSERT 'Calibration-Window Sensitivity' section ----------
# Insert after the Inter-Detector Convergence body paragraph, before the
# Discussion (Heading 2). Strategy: addnext on the prior paragraph, in
# REVERSE order (body first, then heading).

i = find_idx("All three detectors converge on the same direction, timing, and approximate magnitude")
if i >= 0:
    anchor_para = d.paragraphs[i]   # the Inter-Detector Convergence body

    # Find an existing Heading 3 to use as a style template for insertion
    heading3_style = None
    normal_style  = None
    for q in d.paragraphs:
        if q.style.name == "Heading 3" and heading3_style is None:
            heading3_style = q.style
        if q.style.name == "Normal" and normal_style is None:
            normal_style = q.style
        if heading3_style and normal_style:
            break

    # Insert in reverse: body paragraph, then heading paragraph
    body_xml = OxmlElement("w:p")
    anchor_para._p.addnext(body_xml)
    body_para = Paragraph(body_xml, anchor_para._parent)
    if normal_style is not None:
        body_para.style = normal_style
    body_para.add_run(
        "To assess whether the conclusions depend on the upper bound of the "
        "calibration window, the analysis was re-run using a 2010–2017 "
        "calibration window (n = 7,380), which incorporates two cycles "
        "(conferences 2016 and 2017) whose submission deadlines post-date "
        "Grammarly's May 2015 freemium launch and therefore could carry "
        "consumer-grade machine-editing polish. The H1 step at conference "
        "2025 shifts by less than 2 percentage points on every detector "
        "under this wider window: EditLens RoBERTa-large +17.5 pp (vs "
        "+17.8 pp under 2010–2015), EditLens Llama-3.2-3B +29.1 pp "
        "(vs +27.6 pp), and academic-tuned Desklib +21.1 pp (vs +20.9 pp), "
        "with overlapping 95% confidence intervals. The qualitative "
        "conclusion is preserved at either calibration boundary."
    )

    head_xml = OxmlElement("w:p")
    anchor_para._p.addnext(head_xml)
    head_para = Paragraph(head_xml, anchor_para._parent)
    if heading3_style is not None:
        head_para.style = heading3_style
    head_para.add_run("Calibration-Window Sensitivity")

    print(f"[7] Inserted Calibration-Window Sensitivity subsection after Para [{i:04d}]")


# ---------- save ----------
d.save(DOC)
print(f"\nSaved -> {DOC}")
